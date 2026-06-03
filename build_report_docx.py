"""Build a polished Word report for the employee attrition project."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.shared import OxmlElement


OUTPUT = "employee_attrition_analysis_report.docx"


def set_cell_shading(cell, fill):
    # Apply a background fill color to a table cell.
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    # Give table cells comfortable internal padding so text is not cramped.
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    # Mark the row as a header row so Word can repeat it on new pages.
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table):
    # Draw a quiet, consistent grid around the table.
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tbl_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tbl_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "D9DDE3")


def set_table_width_and_indent(table, widths_dxa, indent_dxa=120):
    # Force fixed table geometry so the layout stays stable in Word.
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for w in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(w))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, size=11, color="000000", align=WD_ALIGN_PARAGRAPH.LEFT):
    # Replace the cell content with a single well-formatted paragraph.
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


# Create the document and define the page setup first.
doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]

# Set the body text style so the whole document looks consistent.
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    # Apply consistent typography to the main heading levels.
    s = styles[name]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)

title = doc.add_paragraph()
# The title is built as normal text rather than using Word's Title style.
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(3)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("Employee Attrition Prediction Report")
run.bold = False
run.font.name = "Calibri"
run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
run.font.size = Pt(22)
run.font.color.rgb = RGBColor.from_string("000000")

subtitle = doc.add_paragraph()
# Add a short subtitle to clarify the scope of the report.
subtitle.paragraph_format.space_before = Pt(0)
subtitle.paragraph_format.space_after = Pt(12)
subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
sr = subtitle.add_run("Baseline analysis of the IBM HR employee attrition dataset")
sr.font.name = "Calibri"
sr._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
sr._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
sr.font.size = Pt(11)
sr.italic = True
sr.font.color.rgb = RGBColor.from_string("555555")

meta = doc.add_paragraph()
# Add a small metadata line under the title for context.
meta.paragraph_format.space_before = Pt(0)
meta.paragraph_format.space_after = Pt(12)
meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
mr = meta.add_run("Prepared from the notebook analysis in this project workspace.")
mr.font.name = "Calibri"
mr._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
mr._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
mr.font.size = Pt(10.5)
mr.font.color.rgb = RGBColor.from_string("555555")

doc.add_heading("Executive Summary", level=1)
# This opening paragraph explains the project in plain English.
doc.add_paragraph(
    "This project analyzes employee attrition using the IBM HR dataset and trains a random forest model as a baseline predictor. "
    "The dataset is clean and complete, but the target is imbalanced, so accuracy alone is not enough to judge usefulness. "
    "The current model is better at identifying employees who stay than employees who leave, which is a major limitation for HR use."
)

doc.add_heading("Project Snapshot", level=1)
# A compact summary table makes the report easier to scan quickly.
table = doc.add_table(rows=0, cols=2)
table.style = "Table Grid"
rows = [
    ("Dataset", "WA_Fn-UseC_-HR-Employee-Attrition.csv"),
    ("Rows", "1,470"),
    ("Columns", "35"),
    ("Target", "Attrition"),
    ("Target balance", "No: 1,233; Yes: 237"),
    ("Reported notebook accuracy", "0.8163"),
    ("Re-evaluated accuracy", "0.8435"),
    ("ROC-AUC", "0.7704"),
]
for key, value in rows:
    # Each row pairs a short label with a concise value.
    row = table.add_row()
    set_cell_text(row.cells[0], key, bold=True, size=11, color="1F1F1F")
    set_cell_text(row.cells[1], value, size=11, color="1F1F1F")
    set_cell_shading(row.cells[0], "F2F4F7")
    set_cell_margins(row.cells[0])
    set_cell_margins(row.cells[1])
set_table_width_and_indent(table, [2400, 6960], indent_dxa=120)
set_table_borders(table)

doc.add_heading("Dataset Overview", level=1)
# Explain the data quality and target balance before discussing models.
doc.add_paragraph(
    "The dataset contains 1,470 employee records and 35 features. There are no missing values. "
    "The target distribution is imbalanced, with about 16.1% attrition and 83.9% non-attrition."
)

doc.add_heading("Notebook Workflow", level=1)
# These bullets walk the reader through the original notebook steps.
steps = [
    "Loaded the CSV into pandas.",
    "Converted binary columns such as Attrition, Gender, Over18, and OverTime into numeric form.",
    "One-hot encoded BusinessTravel, Department, EducationField, JobRole, and MaritalStatus.",
    "Dropped EmployeeNumber, EmployeeCount, Over18, and StandardHours.",
    "Visualized distributions with histograms.",
    "Trained a RandomForestClassifier and examined feature importance.",
]
for s in steps:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(s)

doc.add_heading("Model Evaluation", level=1)
# Compare the notebook score to a naive benchmark.
doc.add_paragraph(
    "The notebook reports a test accuracy of 0.8163. That is not much better than the majority-class baseline, which is about 0.8401, "
    "so accuracy is not a strong measure of value for this problem."
)

doc.add_heading("Observed Performance With a Reproducible Split", level=2)
# Show the metrics that matter most for an imbalanced problem.
perf = doc.add_table(rows=0, cols=2)
perf.style = "Table Grid"
perf_rows = [
    ("Accuracy", "0.8435"),
    ("ROC-AUC", "0.7704"),
    ("True negatives", "244"),
    ("False positives", "3"),
    ("False negatives", "43"),
    ("True positives", "4"),
    ("Attrition recall", "0.085"),
    ("Attrition F1-score", "0.148"),
]
for key, value in perf_rows:
    # The left column is the metric name, the right column is the value.
    row = perf.add_row()
    set_cell_text(row.cells[0], key, bold=True, size=11, color="1F1F1F")
    set_cell_text(row.cells[1], value, size=11, color="1F1F1F")
    set_cell_shading(row.cells[0], "F2F4F7")
    set_cell_margins(row.cells[0])
    set_cell_margins(row.cells[1])
set_table_width_and_indent(perf, [3000, 6360], indent_dxa=120)
set_table_borders(perf)

doc.add_paragraph(
    # The key business takeaway is that attrition recall matters more than raw accuracy.
    "The model is highly conservative. It correctly identifies most employees who remain, but it misses most employees who leave. "
    "That makes it weak for proactive attrition prevention."
)

doc.add_heading("Most Influential Features", level=1)
# Feature ranking gives a quick explanation of the strongest model signals.
importance = [
    "MonthlyIncome",
    "Age",
    "TotalWorkingYears",
    "DailyRate",
    "HourlyRate",
    "MonthlyRate",
    "DistanceFromHome",
    "OverTime",
    "YearsWithCurrManager",
    "YearsAtCompany",
]
for item in importance:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

doc.add_heading("Interpretation", level=1)
# Translate the feature ranking into a business explanation.
doc.add_paragraph(
    "The strongest signals are related to pay, career stage, tenure, commuting distance, overtime, and time in role. "
    "These are plausible attrition drivers and line up with common HR patterns."
)

doc.add_heading("Strengths", level=1)
# Keep the project honest by documenting what already works well.
strengths = [
    "Uses a real business dataset with a meaningful target.",
    "Applies appropriate encoding for categorical variables.",
    "Provides a working baseline model and feature-importance view.",
    "Keeps the workflow simple enough to understand quickly.",
]
for item in strengths:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

doc.add_heading("Limitations", level=1)
# Call out the gaps so the project can be improved later.
limitations = [
    "Accuracy is overemphasized even though the classes are imbalanced.",
    "No fixed random state is used, so results are not reproducible.",
    "No stratification is used during train-test splitting.",
    "No cross-validation or hyperparameter tuning is included.",
    "Recall for the attrition class is very low.",
]
for item in limitations:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

doc.add_heading("Recommended Next Steps", level=1)
# These are the practical improvements that would make the project portfolio-ready.
next_steps = [
    "Use a reproducible stratified split with a fixed random_state.",
    "Report precision, recall, F1-score, ROC-AUC, and the confusion matrix.",
    "Try class balancing approaches such as class_weight='balanced' or SMOTE.",
    "Compare multiple models, including logistic regression and gradient boosting.",
    "Tune hyperparameters with cross-validation.",
    "Add explainability with permutation importance or SHAP.",
]
for item in next_steps:
    p = doc.add_paragraph(style="List Number")
    p.add_run(item)

doc.add_heading("Conclusion", level=1)
# Close with a concise summary of the project's current maturity.
doc.add_paragraph(
    "This project is a strong first-pass attrition analysis, but it is not yet a production-ready predictor. "
    "The data contains useful signals, especially around income, age, tenure, and overtime, but the model needs better evaluation and stronger class-imbalance handling before it can reliably support HR decision-making."
)

doc.save(OUTPUT)
print(OUTPUT)
